"""Выгрузка результата в разные форматы.

Все форматы строятся из одной структуры сегментов, поэтому повторное
распознавание для смены формата не требуется.
"""
from __future__ import annotations

import csv
import io
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from ..logging_setup import get_logger
from .postprocess import build_paragraphs

log = get_logger("export")

FORMATS = ["txt", "json", "srt", "vtt", "ass", "tsv", "csv", "md", "docx"]


def format_timestamp(seconds: float, style: str = "srt") -> str:
    # Всё считаем от округлённых миллисекунд. Раньше округление добавляло
    # секунду отдельно, и перенос дальше не шёл: 59.9996 давало
    # «00:00:60,000», а 3599.9996 — «00:59:60,000». Строгие плееры такой
    # блок субтитров просто отбрасывают.
    total_ms = int(round(max(0.0, float(seconds)) * 1000))
    millis = total_ms % 1000
    total_seconds = total_ms // 1000
    secs = total_seconds % 60
    minutes = (total_seconds // 60) % 60
    hours = total_seconds // 3600
    if style == "srt":
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
    if style == "vtt":
        return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"
    if style == "ass":
        return f"{hours:d}:{minutes:02d}:{secs:02d}.{millis // 10:02d}"
    if style == "seconds":
        return f"{total_ms / 1000:.3f}"
    return f"{hours:02d}:{minutes:02d}:{secs:02d}.{millis:03d}"


def wrap_subtitle(text: str, max_width: int = 42, max_lines: int = 2) -> str:
    """Аккуратно разбивает реплику на строки субтитра по границам слов.

    Слова длиннее строки (ссылки, длинные числа, склеенные термины)
    разрываются принудительно — иначе строка вылезет за пределы кадра.
    """
    raw_words = text.split()
    words: list[str] = []
    for word in raw_words:
        while len(word) > max_width:
            words.append(word[:max_width])
            word = word[max_width:]
        if word:
            words.append(word)
    if not words:
        return ""
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if len(candidate) <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    if len(lines) > max_lines:
        # Перераспределяем равномерно, чтобы не оставлять «висячее» слово.
        # Ширину подбираем, пока текст не уложится: прежний расчёт «на глаз»
        # (длина делить на число строк) не гарантировал нужного числа строк,
        # а лишние строки просто отрезались — вместе со словами. На реплике
        # из пятнадцати слов одно пропадало без следа.
        joined = " ".join(words)
        target = max(1, len(joined) // max_lines + 1)
        best = lines
        for _ in range(24):
            packed, current = [], ""
            for word in words:
                candidate = f"{current} {word}".strip()
                if len(candidate) <= target or not current:
                    current = candidate
                else:
                    packed.append(current)
                    current = word
            if current:
                packed.append(current)
            best = packed
            if len(packed) <= max_lines:
                break
            target += max(1, target // 8)
        # Даже если ширину пришлось увеличить сверх max_width, текст остаётся
        # целым: обрезанная реплика хуже длинной строки.
        lines = best
    return "\n".join(lines)


def _speaker_prefix(seg: dict[str, Any], settings: dict[str, Any]) -> str:
    if not settings.get("include_speaker_labels", True):
        return ""
    speaker = seg.get("speaker")
    return f"{speaker}: " if speaker else ""


# --------------------------------------------------------------------------
# Текстовые форматы
# --------------------------------------------------------------------------

def to_txt(result: dict[str, Any], settings: dict[str, Any]) -> str:
    segments = result.get("segments", [])
    mode = str(settings.get("paragraph_mode") or "speaker")
    paragraphs = build_paragraphs(
        segments, mode,
        speaker_labels=bool(settings.get("include_speaker_labels", True)))
    return "\n\n".join(p for p in paragraphs if p.strip()) + "\n"


def to_markdown(result: dict[str, Any], settings: dict[str, Any]) -> str:
    meta = result.get("meta", {})
    lines = [f"# Расшифровка: {meta.get('filename', 'запись')}", ""]
    lines.append(f"- **Модель:** {meta.get('model', '')}")
    lines.append(f"- **Язык:** {meta.get('language', '')}")
    dur = meta.get("duration_s")
    if dur:
        lines.append(f"- **Длительность:** {format_timestamp(dur, 'hms')}")
    if meta.get("created_at"):
        lines.append(f"- **Обработано:** {meta['created_at']}")
    if result.get("summary"):
        lines += ["", "## Краткое содержание", "", result["summary"]]
    lines += ["", "## Текст", ""]
    for seg in result.get("segments", []):
        stamp = format_timestamp(seg.get("start", 0), "hms")
        prefix = _speaker_prefix(seg, settings)
        lines.append(f"**[{stamp}]** {prefix}{seg.get('text', '')}")
        lines.append("")
    return "\n".join(lines)


def to_json(result: dict[str, Any], settings: dict[str, Any]) -> str:
    payload = dict(result)
    if not settings.get("include_confidence", True):
        for seg in payload.get("segments", []):
            seg.pop("confidence", None)
            for word in seg.get("words") or []:
                word.pop("confidence", None)
    if not settings.get("word_timestamps", True):
        for seg in payload.get("segments", []):
            seg.pop("words", None)
    return json.dumps(payload, ensure_ascii=False, indent=2)


# --------------------------------------------------------------------------
# Субтитры
# --------------------------------------------------------------------------

def to_srt(result: dict[str, Any], settings: dict[str, Any]) -> str:
    width = int(settings.get("subtitle_max_line_width") or 42)
    lines_limit = int(settings.get("subtitle_max_lines") or 2)
    min_dur = float(settings.get("subtitle_min_duration_s") or 1.0)
    out: list[str] = []
    for idx, seg in enumerate(_prepare_subtitles(result, min_dur), start=1):
        text = wrap_subtitle(_speaker_prefix(seg, settings) + seg.get("text", ""),
                             width, lines_limit)
        if not text.strip():
            continue
        out.append(str(idx))
        out.append(f"{format_timestamp(seg['start'], 'srt')} --> "
                   f"{format_timestamp(seg['end'], 'srt')}")
        out.append(text)
        out.append("")
    return "\n".join(out)


def to_vtt(result: dict[str, Any], settings: dict[str, Any]) -> str:
    width = int(settings.get("subtitle_max_line_width") or 42)
    lines_limit = int(settings.get("subtitle_max_lines") or 2)
    min_dur = float(settings.get("subtitle_min_duration_s") or 1.0)
    out = ["WEBVTT", ""]
    for seg in _prepare_subtitles(result, min_dur):
        text = wrap_subtitle(_speaker_prefix(seg, settings) + seg.get("text", ""),
                             width, lines_limit)
        if not text.strip():
            continue
        out.append(f"{format_timestamp(seg['start'], 'vtt')} --> "
                   f"{format_timestamp(seg['end'], 'vtt')}")
        out.append(text)
        out.append("")
    return "\n".join(out)


def _ass_escape(text: str) -> str:
    """Обезвреживает разметку ASS внутри реплики.

    В фигурных скобках libass ждёт команды оформления, поэтому «скидка {30}
    процентов» отрисовывалась без числа: всё, что в скобках, считалось
    командой и просто не показывалось. Экранируем скобки и обратный слэш,
    который в ASS тоже управляющий.
    """
    return (text.replace("\\", "\\\\")
                .replace("{", "\\{")
                .replace("}", "\\}"))


def to_ass(result: dict[str, Any], settings: dict[str, Any]) -> str:
    width = int(settings.get("subtitle_max_line_width") or 42)
    lines_limit = int(settings.get("subtitle_max_lines") or 2)
    min_dur = float(settings.get("subtitle_min_duration_s") or 1.0)
    header = [
        "[Script Info]",
        "ScriptType: v4.00+",
        "PlayResX: 1920",
        "PlayResY: 1080",
        "WrapStyle: 0",
        "ScaledBorderAndShadow: yes",
        "",
        "[V4+ Styles]",
        ("Format: Name, Fontname, Fontsize, PrimaryColour, SecondaryColour, OutlineColour, "
         "BackColour, Bold, Italic, Underline, StrikeOut, ScaleX, ScaleY, Spacing, Angle, "
         "BorderStyle, Outline, Shadow, Alignment, MarginL, MarginR, MarginV, Encoding"),
        ("Style: Default,Arial,54,&H00FFFFFF,&H000000FF,&H00000000,&H80000000,0,0,0,0,"
         "100,100,0,0,1,3,1,2,80,80,60,204"),
        "",
        "[Events]",
        "Format: Layer, Start, End, Style, Name, MarginL, MarginR, MarginV, Effect, Text",
    ]
    for seg in _prepare_subtitles(result, min_dur):
        text = wrap_subtitle(_ass_escape(seg.get("text", "")),
                             width, lines_limit).replace("\n", "\\N")
        if not text.strip():
            continue
        speaker = seg.get("speaker") or ""
        header.append(f"Dialogue: 0,{format_timestamp(seg['start'], 'ass')},"
                      f"{format_timestamp(seg['end'], 'ass')},Default,{speaker},0,0,0,,{text}")
    return "\n".join(header) + "\n"


def _prepare_subtitles(result: dict[str, Any], min_duration: float) -> list[dict[str, Any]]:
    """Гарантирует минимальную длительность и отсутствие наложений."""
    segments = [dict(s) for s in result.get("segments", []) if s.get("text", "").strip()]
    for idx, seg in enumerate(segments):
        start = float(seg.get("start", 0.0))
        end = float(seg.get("end", start))
        if end - start < min_duration:
            end = start + min_duration
        next_start = float(segments[idx + 1].get("start", end)) if idx + 1 < len(segments) else None
        if next_start is not None and end > next_start:
            end = max(start + 0.2, next_start - 0.04)
        seg["start"], seg["end"] = start, end
    return segments


# --------------------------------------------------------------------------
# Таблицы
# --------------------------------------------------------------------------

def to_table(result: dict[str, Any], settings: dict[str, Any], delimiter: str = "\t") -> str:
    buffer = io.StringIO()
    writer = csv.writer(buffer, delimiter=delimiter, lineterminator="\n",
                        quoting=csv.QUOTE_MINIMAL)
    header = ["№", "начало_с", "конец_с", "длительность_с", "говорящий", "текст"]
    if settings.get("include_confidence", True):
        header.append("уверенность")
    header += ["слов", "символов"]
    writer.writerow(header)
    for idx, seg in enumerate(result.get("segments", []), start=1):
        text = seg.get("text", "")
        row = [idx, f"{seg.get('start', 0):.3f}", f"{seg.get('end', 0):.3f}",
               f"{float(seg.get('end', 0)) - float(seg.get('start', 0)):.3f}",
               seg.get("speaker") or "", text]
        if settings.get("include_confidence", True):
            conf = seg.get("confidence")
            row.append(f"{conf:.4f}" if conf is not None else "")
        row += [len(text.split()), len(text)]
        writer.writerow(row)
    return buffer.getvalue()


# --------------------------------------------------------------------------
# Word
# --------------------------------------------------------------------------

def to_docx(result: dict[str, Any], settings: dict[str, Any], path: Path) -> Path:
    """Создаёт документ Word. При отсутствии python-docx выгружает Markdown."""
    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
    except ModuleNotFoundError:
        fallback = path.with_suffix(".md")
        fallback.write_text(to_markdown(result, settings), encoding="utf-8")
        log.warning("python-docx не установлен, документ сохранён как Markdown: %s", fallback)
        return fallback

    meta = result.get("meta", {})
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(f"Расшифровка: {meta.get('filename', 'запись')}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    info = doc.add_paragraph()
    info.add_run("Модель: ").bold = True
    info.add_run(f"{meta.get('model', '')}    ")
    info.add_run("Язык: ").bold = True
    info.add_run(f"{meta.get('language', '')}    ")
    if meta.get("duration_s"):
        info.add_run("Длительность: ").bold = True
        info.add_run(format_timestamp(meta["duration_s"], "hms"))
    stamp = doc.add_paragraph()
    stamp.add_run(f"Обработано: {meta.get('created_at', datetime.now().strftime('%d.%m.%Y %H:%M'))}")
    for run in stamp.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    if result.get("summary"):
        doc.add_heading("Краткое содержание", level=1)
        doc.add_paragraph(result["summary"])

    doc.add_heading("Текст", level=1)
    prev_speaker = object()
    for seg in result.get("segments", []):
        text = seg.get("text", "").strip()
        if not text:
            continue
        speaker = seg.get("speaker")
        para = doc.add_paragraph()
        if settings.get("include_speaker_labels", True) and speaker and speaker != prev_speaker:
            para.add_run(f"{speaker}").bold = True
            para.add_run(f"  [{format_timestamp(seg.get('start', 0), 'hms')}]").font.size = Pt(8)
            para = doc.add_paragraph()
        elif not speaker:
            mark = para.add_run(f"[{format_timestamp(seg.get('start', 0), 'hms')}] ")
            mark.font.size = Pt(8)
            mark.font.color.rgb = RGBColor(0x90, 0x90, 0x90)
        para.add_run(text)
        prev_speaker = speaker

    metrics = result.get("metrics") or {}
    if metrics:
        doc.add_page_break()
        doc.add_heading("Показатели обработки", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        table.rows[0].cells[0].text = "Показатель"
        table.rows[0].cells[1].text = "Значение"
        labels = {
            "rtf": "Коэффициент реального времени (RTF)",
            "processing_time_s": "Время обработки, с",
            "segments": "Сегментов",
            "words": "Слов",
            "avg_confidence": "Средняя уверенность",
            "speech_ratio": "Доля речи",
        }
        for key, label in labels.items():
            if key in metrics and metrics[key] is not None:
                row = table.add_row()
                row.cells[0].text = label
                value = metrics[key]
                row.cells[1].text = (f"{value:.3f}" if isinstance(value, float) else str(value))
    doc.save(str(path))
    return path


# --------------------------------------------------------------------------
# Точка входа
# --------------------------------------------------------------------------

def write_all(result: dict[str, Any], settings: dict[str, Any], outdir: Path,
              basename: str) -> dict[str, str]:
    """Сохраняет все запрошенные форматы. Возвращает соответствие формат → путь."""
    outdir.mkdir(parents=True, exist_ok=True)
    formats = settings.get("output_formats") or ["txt", "json"]
    if isinstance(formats, str):
        formats = [formats]
    safe = re.sub(r"[^\w\-. ]+", "_", basename).strip() or "result"
    written: dict[str, str] = {}

    handlers = {
        "txt": lambda p: p.write_text(to_txt(result, settings), encoding="utf-8"),
        "md": lambda p: p.write_text(to_markdown(result, settings), encoding="utf-8"),
        "json": lambda p: p.write_text(to_json(result, settings), encoding="utf-8"),
        "srt": lambda p: p.write_text(to_srt(result, settings), encoding="utf-8"),
        "vtt": lambda p: p.write_text(to_vtt(result, settings), encoding="utf-8"),
        "ass": lambda p: p.write_text(to_ass(result, settings), encoding="utf-8"),
        "tsv": lambda p: p.write_text(to_table(result, settings, "\t"), encoding="utf-8"),
        "csv": lambda p: p.write_text(to_table(result, settings, ","), encoding="utf-8-sig"),
    }

    for fmt in formats:
        try:
            if fmt == "docx":
                path = to_docx(result, settings, outdir / f"{safe}.docx")
                written["docx"] = str(path)
                continue
            handler = handlers.get(fmt)
            if handler is None:
                log.warning("Неизвестный формат выгрузки: %s", fmt)
                continue
            path = outdir / f"{safe}.{fmt}"
            handler(path)
            written[fmt] = str(path)
        except Exception as exc:
            log.error("Не удалось сохранить формат %s: %s", fmt, exc)
    return written
